from typing import Any, Dict, Optional
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from AkvoFormPrint.parsers.akvo_flow_parser import AkvoFlowFormParser
from AkvoFormPrint.parsers.akvo_arf_parser import AkvoReactFormParser
from AkvoFormPrint.parsers.default_parser import DefaultParser
from AkvoFormPrint.parsers.base_parser import BaseParser
from AkvoFormPrint.enums import HintText, QuestionType


class DocxRenderer:
    def __init__(
        self,
        orientation: str = "landscape",
        add_section_numbering: bool = False,
        add_question_numbering: bool = True,
        parser_type: Optional[str] = None,
        raw_json: Optional[Dict[str, Any]] = None,
        output_path: Optional[str] = "output.docx",
    ):
        """Initialize DocxRenderer with configuration."""
        assert orientation in (
            "portrait",
            "landscape",
        ), "Orientation must be 'portrait' or 'landscape'"
        self.orientation = orientation
        self.add_section_numbering = add_section_numbering
        self.add_question_numbering = add_question_numbering
        self.parser_type = parser_type
        self.raw_json = raw_json
        self.output_path = output_path

        self.parser: Optional[BaseParser] = None
        self.form_model = None
        self.doc = Document()

        if parser_type or raw_json:
            self.parser = self._get_parser(parser_type)

        if self.raw_json:
            self.form_model = self._parse_form()

    def _get_parser(self, parser_type: Optional[str] = None) -> BaseParser:
        if parser_type == "flow":
            return AkvoFlowFormParser()
        elif parser_type == "arf":
            return AkvoReactFormParser()
        elif parser_type is None or parser_type == "default":
            return DefaultParser()
        else:
            raise ValueError(f"Unknown parser type: {parser_type}")

    def _parse_form(self):
        if not self.raw_json:
            raise ValueError("No raw_json data provided to parse")
        if not self.parser:
            self.parser = self._get_parser(self.parser_type)
        form_model = self.parser.parse(self.raw_json)
        return self._inject_question_numbers(form_model)

    def _inject_question_numbers(self, form):
        section_index = 0
        counter = 1
        for section in form.sections:
            section.letter = (
                self._number_to_letter(section_index)
                if self.add_section_numbering
                else None
            )
            if self.add_section_numbering:
                section_index += 1
            for question in section.questions:
                question.number = (
                    counter if self.add_question_numbering else None
                )
                if self.add_question_numbering:
                    counter += 1
        return form

    def _number_to_letter(self, n: int) -> str:
        """Convert a zero-based index to letters (A, B, ..., Z, AA, AB, ...)."""
        result = ""
        while n >= 0:
            result = chr(n % 26 + ord("A")) + result
            n = n // 26 - 1
        return result

    def _set_landscape(self):
        """Set the document orientation to landscape with specific margins."""
        section = self.doc.sections[-1]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = (
            section.page_height,
            section.page_width,
        )
        for side in [
            "top_margin",
            "bottom_margin",
            "left_margin",
            "right_margin",
        ]:
            setattr(section, side, Inches(0.5))

    def _insert_continuous_section_break(self, paragraph):
        sectPr = OxmlElement("w:sectPr")
        type_elem = OxmlElement("w:type")
        type_elem.set(qn("w:val"), "continuous")
        sectPr.append(type_elem)

        p = paragraph._p
        p.addnext(sectPr)

    def _set_paragraph_shading_and_underline(
        self,
        paragraph,
        shading: Optional[bool] = True,
        shading_color: Optional[str] = "#F2F2F2",
        underline: Optional[bool] = False,
        underline_color: Optional[str] = "000000",
    ):
        p = paragraph._p
        pPr = p.get_or_add_pPr()

        # Shading
        if shading:
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), shading_color)
            pPr.append(shd)

        # Underline (bottom border)
        if underline:
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")  # 0.5pt line
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), underline_color)
            pBdr.append(bottom)
            pPr.append(pBdr)

    def _add_horizontal_line(self):
        para = self.doc.add_paragraph()
        p = para._p
        pPr = p.get_or_add_pPr()

        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")  # 0.5pt line
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")

        pBdr.append(bottom)
        pPr.append(pBdr)

    def render_docx(self):
        """Render the DOCX document with the parsed form model."""
        if not self.form_model:
            self.form_model = self._parse_form()

        if self.orientation == "landscape":
            self._set_landscape()

        # Title
        title_para = self.doc.add_paragraph(
            self.form_model.title, style="Title"
        )
        title_para.style.font.size = Pt(14)

        # Insert a continuous section break directly after the title
        self._insert_continuous_section_break(paragraph=title_para)

        # Set the 2-column layout for the new section (after the title)
        section = self.doc.sections[-1]
        sectPr = section._sectPr
        cols_elem = sectPr.xpath("./w:cols")
        cols = cols_elem[0] if cols_elem else OxmlElement("w:cols")
        cols.set(qn("w:num"), "2")
        if not cols_elem:
            sectPr.append(cols)

        # Sections and questions
        for idx, section_data in enumerate(self.form_model.sections):
            if idx != 0:
                self.doc.add_page_break()

            # Section title
            section_title = (
                f"{section_data.letter}. {section_data.title}"
                if section_data.letter
                else section_data.title
            )
            section_para = self.doc.add_paragraph(
                section_title, style="Heading 1"
            )
            section_para.style.font.size = Pt(12)
            self._set_paragraph_shading_and_underline(
                paragraph=section_para,
                shading_color="#D9D9D9",
                underline=True,
            )

            # Questions
            for qidx, question in enumerate(section_data.questions):
                # Question
                required_mark = "*" if question.answer.required else ""
                qtext = (
                    f"{question.number}. {question.label} {required_mark}"
                    if question.number
                    else question.label
                )
                question_para = self.doc.add_paragraph(qtext)
                self._set_paragraph_shading_and_underline(
                    paragraph=question_para
                )
                question_para.style.font.size = Pt(10)
                question_para.paragraph_format.space_before = Pt(10)
                question_para.paragraph_format.space_after = (
                    Pt(5) if not question.hint else Pt(2)
                )

                # If hint is present, add as a sub-paragraph
                if question.hint:
                    hint_para = self.doc.add_paragraph()
                    # Add the hint text as a run so we can set
                    # italic + size only for the hint
                    hint_run = hint_para.add_run(question.hint)
                    hint_run.italic = True
                    hint_run.font.size = Pt(9)
                    hint_para.paragraph_format.space_before = Pt(0)
                    hint_para.paragraph_format.space_after = Pt(5)
                    # Add shading to hint
                    self._set_paragraph_shading_and_underline(
                        paragraph=hint_para,
                    )

                # Dependency hint
                if question.dependencies:
                    dependency_text = "Answer only if "
                    for idx, dep in enumerate(question.dependencies):
                        dependency_text += f'"{dep.expected_answer}" selected '
                        dependency_text += (
                            f'for question "{dep.depends_on_question_id}"'
                        )
                        if idx != len(question.dependencies) - 1:
                            dependency_text += " AND "
                    dependency_para = self.doc.add_paragraph()
                    # Add the hint text as a run so we can set
                    # italic + size only for the hint
                    dependency_run = dependency_para.add_run(dependency_text)
                    dependency_run.italic = True
                    dependency_run.font.size = Pt(9)
                    hint_para.paragraph_format.space_before = Pt(0)
                    hint_para.paragraph_format.space_after = Pt(2)

                if question.type in [
                    QuestionType.OPTION,
                    QuestionType.MULTIPLE_OPTION,
                ]:
                    # option symbol
                    option_symbol = (
                        "( )"
                        if question.type == QuestionType.OPTION
                        else "[ ]"
                    )
                    # Create a table directly in the document
                    table = self.doc.add_table(rows=1, cols=2)
                    table.autofit = True

                    # Fill columns with options, reusing first empty paragraph
                    col1Len = 9 if qidx == 0 else 5
                    col1 = question.answer.options[:col1Len]
                    col2 = question.answer.options[col1Len:]

                    # First column
                    cell1 = table.cell(0, 0)
                    if col1:
                        first_para = cell1.paragraphs[0]
                        first_para.text = f"{option_symbol} {col1[0]}"
                        first_para.style.font.size = Pt(10)
                        first_para.paragraph_format.space_after = Pt(0)
                        first_para.paragraph_format.space_before = Pt(0)

                        for opt in col1[1:]:
                            para = cell1.add_paragraph(
                                f"{option_symbol} {opt}"
                            )
                            para.style.font.size = Pt(10)
                            para.paragraph_format.space_after = Pt(0)
                            para.paragraph_format.space_before = Pt(0)

                    # Second column
                    cell2 = table.cell(0, 1)
                    if col2:
                        first_para = cell2.paragraphs[0]
                        first_para.text = f"{option_symbol} {col2[0]}"
                        first_para.style.font.size = Pt(10)
                        first_para.paragraph_format.space_after = Pt(0)
                        first_para.paragraph_format.space_before = Pt(0)

                        for opt in col2[1:]:
                            para = cell2.add_paragraph(
                                f"{option_symbol} {opt}"
                            )
                            para.style.font.size = Pt(10)
                            para.paragraph_format.space_after = Pt(0)
                            para.paragraph_format.space_before = Pt(0)

                elif question.type == QuestionType.NUMBER:
                    num_boxes = question.answer.numberBox
                    if num_boxes:
                        table = self.doc.add_table(rows=1, cols=num_boxes)
                        table.autofit = False

                        # Indent the table from the left
                        tbl = table._tbl
                        tblPr = tbl.tblPr
                        if tblPr is None:
                            tblPr = OxmlElement("w:tblPr")
                            tbl.insert(0, tblPr)

                        tblInd = OxmlElement("w:tblInd")
                        tblInd.set(
                            qn("w:w"), str(int(Inches(0.15).inches * 1440))
                        )  # 0.25 inch indent
                        tblInd.set(qn("w:type"), "dxa")
                        tblPr.append(tblInd)

                        box_width = Inches(0.2)
                        for col_idx in range(num_boxes):
                            cell = table.cell(0, col_idx)

                            # Set fixed width
                            tcPr = cell._tc.get_or_add_tcPr()
                            tcW = OxmlElement("w:tcW")
                            tcW.set(
                                qn("w:w"), str(int(box_width.inches * 1440))
                            )  # width in twips
                            tcW.set(qn("w:type"), "dxa")
                            tcPr.append(tcW)

                            # Add borders
                            borders = OxmlElement("w:tcBorders")
                            for border_name in [
                                "top",
                                "left",
                                "bottom",
                                "right",
                            ]:
                                border = OxmlElement(f"w:{border_name}")
                                border.set(qn("w:val"), "single")
                                border.set(qn("w:sz"), "4")  # 0.5pt border
                                border.set(qn("w:space"), "0")
                                border.set(qn("w:color"), "000000")
                                borders.append(border)
                            tcPr.append(borders)

                            para = cell.paragraphs[0]
                            para.text = ""
                            para.paragraph_format.space_after = Pt(0)
                            para.paragraph_format.space_before = Pt(0)

                    number_hint_text = None
                    if question.answer.minValue and question.answer.maxValue:
                        number_hint_text = "Enter a number between "
                        number_hint_text += f"{question.answer.minValue} and "
                        number_hint_text += f"{question.answer.maxValue}"
                    elif question.answer.minValue:
                        number_hint_text = "Min: "
                        number_hint_text += f"{question.answer.minValue}"
                    elif question.answer.maxValue:
                        number_hint_text = "Max: "
                        number_hint_text += f"{question.answer.maxValue}"
                    if number_hint_text:
                        hint_para = self.doc.add_paragraph()
                        # Add the hint text as a run so we can set
                        # italic + size only for the hint
                        hint_run = hint_para.add_run(number_hint_text)
                        hint_run.italic = True
                        hint_run.font.size = Pt(9)
                        hint_para.paragraph_format.space_before = Pt(4)
                        hint_para.paragraph_format.space_after = Pt(0)

                elif question.type == QuestionType.DATE:
                    date_para = self.doc.add_paragraph(
                        "[    ][    ] / [    ][    ] / [    ][    ][    ][    ]"
                    )
                    date_para.style.font.size = Pt(14)
                    date_para.paragraph_format.space_after = Pt(0)

                    hint_para = self.doc.add_paragraph()
                    # Add the hint text as a run so we can set
                    # italic + size only for the hint
                    hint_run = hint_para.add_run(HintText.DATE.value)
                    hint_run.italic = True
                    hint_run.font.size = Pt(9)
                    hint_para.paragraph_format.space_before = Pt(2)
                    hint_para.paragraph_format.space_after = Pt(0)

                elif question.type == QuestionType.TEXT:
                    # Add the first horizontal line
                    self._add_horizontal_line()

                    # Optional: add a small gap between the lines
                    spacer_para = self.doc.add_paragraph()
                    spacer_para.paragraph_format.space_after = Pt(0)
                    spacer_para.paragraph_format.space_before = Pt(0)

                    # Add the second horizontal line
                    self._add_horizontal_line()

                else:
                    # INPUT TYPE
                    self._add_horizontal_line()

        self.doc.save(self.output_path)
